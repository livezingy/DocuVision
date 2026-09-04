"""
Export Service - Multi-format document export (JSON, CSV, Markdown, Word)
"""

from typing import Dict, Any, List
from loguru import logger
import os
import json
import csv
import io

from app.core.config import settings


def table_confidence_pct(table: Dict[str, Any]) -> int:
    """Return table confidence as a 0-100 integer.

    Prefers ``confidence`` (layout detector score mapped by layout_service,
    typically 0-1). Falls back to ``score``. Values in [0, 1] are ratios;
    values greater than 1 are treated as already-percent.
    """
    raw = table.get("confidence")
    if raw is None:
        raw = table.get("score")
    if raw is None:
        return 0
    try:
        val = float(raw)
    except (TypeError, ValueError):
        return 0
    if 0.0 <= val <= 1.0:
        return int(round(val * 100.0))
    return int(round(val))


def format_table_export_title(index_1based: int, table: Dict[str, Any]) -> str:
    """Human-readable table title including page and confidence percent."""
    page = table.get("page", "?")
    pct = table_confidence_pct(table)
    return f"Table {index_1based} (Page {page}) confidence={pct}%"


def format_table_csv_banner(index_1based: int, table: Dict[str, Any]) -> str:
    """CSV separator line placed above each table body."""
    return f"=== {format_table_export_title(index_1based, table)} ==="


class ExportService:
    """
    Export Service for multiple output formats
    
    Supported formats:
    - JSON: Full structured data
    - CSV: Table data
    - Markdown: Formatted document
    - Word (DOCX): Microsoft Word document
    - Excel (XLSX): Spreadsheet with tables
    """
    
    def __init__(self):
        self.output_dir = settings.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)
    
    async def to_json(
        self, 
        result: Dict[str, Any], 
        task_id: str,
        pretty: bool = True
    ) -> str:
        """
        Export full result to JSON file
        
        Args:
            result: Processing result dictionary
            task_id: Task identifier
            pretty: Whether to format JSON with indentation
        
        Returns:
            Path to exported JSON file
        """
        output_path = os.path.join(self.output_dir, f"{task_id}_result.json")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            if pretty:
                json.dump(result, f, ensure_ascii=False, indent=2)
            else:
                json.dump(result, f, ensure_ascii=False)
        
        logger.info(f"Exported JSON to {output_path}")
        return output_path
    
    async def to_csv(
        self, 
        result: Dict[str, Any], 
        task_id: str,
        include_all_tables: bool = True
    ) -> str:
        """
        Export tables to CSV file
        
        Args:
            result: Processing result dictionary
            task_id: Task identifier
            include_all_tables: Whether to include all tables in one file
        
        Returns:
            Path to exported CSV file
        """
        output_path = os.path.join(self.output_dir, f"{task_id}_tables.csv")
        
        tables = result.get('tables', [])
        
        with open(output_path, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.writer(f)
            
            for idx, table in enumerate(tables):
                if 'data' in table and table['data']:
                    writer.writerow([format_table_csv_banner(idx + 1, table)])
                    
                    # Write table data
                    for row in table['data']:
                        writer.writerow(row)
                    
                    # Separator between tables
                    writer.writerow([])
        
        logger.info(f"Exported CSV to {output_path}")
        return output_path
    
    async def to_markdown(self, result: Dict[str, Any]) -> str:
        """
        Export result to Markdown format
        
        Args:
            result: Processing result dictionary
        
        Returns:
            Markdown formatted string
        """
        doc_info = result.get('document_info', {})
        layout = result.get('layout', {})
        text_blocks = result.get('text_blocks', [])
        tables = result.get('tables', [])
        keywords = result.get('keywords', [])
        
        md_lines = []
        
        # Document Title
        md_lines.append(f"# {doc_info.get('file_name', 'Document Analysis Result')}")
        md_lines.append("")
        
        # Metadata
        md_lines.append("## 📋 Document Information")
        md_lines.append("")
        md_lines.append(f"| Property | Value |")
        md_lines.append(f"| --- | --- |")
        md_lines.append(f"| **File Name** | {doc_info.get('file_name', 'N/A')} |")
        md_lines.append(f"| **Processed At** | {doc_info.get('processed_at', 'N/A')} |")
        md_lines.append(f"| **Total Pages** | {doc_info.get('pages', 'N/A')} |")
        
        if 'engine_used' in result:
            md_lines.append(f"| **OCR Engine** | {result.get('engine_used', 'N/A')} |")
        
        md_lines.append("")
        
        # Document Structure Summary
        if layout:
            summary = layout.get('summary', {})
            type_counts = summary.get('type_counts', {})
            
            md_lines.append("## 📊 Document Structure")
            md_lines.append("")
            md_lines.append(f"Total Elements: **{summary.get('total_elements', 0)}**")
            md_lines.append("")
            
            if type_counts:
                md_lines.append("| Element Type | Count |")
                md_lines.append("| --- | --- |")
                for elem_type, count in type_counts.items():
                    md_lines.append(f"| {elem_type.title()} | {count} |")
                md_lines.append("")
        
        # Extracted Text
        if text_blocks:
            md_lines.append("## 📝 Extracted Text")
            md_lines.append("")
            
            current_page = 0
            for block in text_blocks:
                page = block.get('page', 1)
                if page != current_page:
                    md_lines.append(f"### Page {page}")
                    md_lines.append("")
                    current_page = page
                
                text = block.get('text', '')
                confidence = block.get('confidence', 0)
                
                # Format based on confidence
                if confidence >= 0.95:
                    md_lines.append(f"{text}")
                else:
                    md_lines.append(f"{text} _(confidence: {confidence:.1%})_")
                md_lines.append("")
        
        # Tables
        if tables:
            md_lines.append("## 📋 Extracted Tables")
            md_lines.append("")
            
            for idx, table in enumerate(tables):
                md_lines.append(f"### {format_table_export_title(idx + 1, table)}")
                md_lines.append(f"_{table.get('rows', '?')} rows × {table.get('columns', '?')} columns_")
                md_lines.append("")
                
                data = table.get('data', [])
                if data:
                    # Markdown table header
                    header = data[0]
                    md_lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
                    md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                    
                    # Data rows
                    for row in data[1:]:
                        # Ensure row has same number of columns as header
                        padded_row = row + [''] * (len(header) - len(row)) if len(row) < len(header) else row[:len(header)]
                        md_lines.append("| " + " | ".join(str(cell) for cell in padded_row) + " |")
                    
                    md_lines.append("")
        
        # Keywords
        if keywords:
            md_lines.append("## 🏷️ Keywords")
            md_lines.append("")
            md_lines.append(", ".join(f"`{kw}`" for kw in keywords))
            md_lines.append("")
        
        # Footer
        md_lines.append("---")
        md_lines.append("_Generated by DocuVision - Intelligent Document Processing System_")
        
        return "\n".join(md_lines)
    
    async def to_markdown_file(self, result: Dict[str, Any], task_id: str) -> str:
        """Export to Markdown file"""
        output_path = os.path.join(self.output_dir, f"{task_id}_result.md")
        
        md_content = await self.to_markdown(result)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        logger.info(f"Exported Markdown to {output_path}")
        return output_path
    
    async def to_docx(self, result: Dict[str, Any], task_id: str) -> str:
        """
        Export result to Word document
        
        Args:
            result: Processing result dictionary
            task_id: Task identifier
        
        Returns:
            Path to exported DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            logger.error("python-docx not installed, cannot export Word document")
            raise RuntimeError("python-docx not installed")
        
        output_path = os.path.join(self.output_dir, f"{task_id}_result.docx")
        
        doc = Document()
        
        doc_info = result.get('document_info', {})
        text_blocks = result.get('text_blocks', [])
        tables = result.get('tables', [])
        keywords = result.get('keywords', [])
        layout = result.get('layout', {})
        
        # Title
        title = doc.add_heading(doc_info.get('file_name', 'Document Analysis Result'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Document Info Section
        doc.add_heading('Document Information', level=1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('File Name', doc_info.get('file_name', 'N/A')),
            ('Processed At', doc_info.get('processed_at', 'N/A')),
            ('Total Pages', str(doc_info.get('pages', 'N/A'))),
            ('OCR Engine', result.get('engine_used', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Document Structure Summary
        if layout:
            doc.add_heading('Document Structure', level=1)
            
            summary = layout.get('summary', {})
            type_counts = summary.get('type_counts', {})
            
            if type_counts:
                struct_table = doc.add_table(rows=len(type_counts) + 1, cols=2)
                struct_table.style = 'Table Grid'
                
                struct_table.rows[0].cells[0].text = 'Element Type'
                struct_table.rows[0].cells[1].text = 'Count'
                
                for i, (elem_type, count) in enumerate(type_counts.items(), 1):
                    struct_table.rows[i].cells[0].text = elem_type.title()
                    struct_table.rows[i].cells[1].text = str(count)
            
            doc.add_paragraph()
        
        # Extracted Text
        if text_blocks:
            doc.add_heading('Extracted Text', level=1)
            
            current_page = 0
            for block in text_blocks:
                page = block.get('page', 1)
                if page != current_page:
                    doc.add_heading(f'Page {page}', level=2)
                    current_page = page
                
                text = block.get('text', '')
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)
        
        # Tables
        if tables:
            doc.add_heading('Extracted Tables', level=1)
            
            for idx, table_data in enumerate(tables):
                doc.add_heading(format_table_export_title(idx + 1, table_data), level=2)
                
                meta_p = doc.add_paragraph()
                meta_p.add_run(f'{table_data.get("rows", "?")} rows × {table_data.get("columns", "?")} columns')
                
                data = table_data.get('data', [])
                if data:
                    rows = len(data)
                    cols = len(data[0]) if data else 0
                    
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        for i, row_data in enumerate(data):
                            row = table.rows[i]
                            for j, cell_text in enumerate(row_data):
                                if j < len(row.cells):
                                    row.cells[j].text = str(cell_text)
                
                doc.add_paragraph()
        
        # Keywords
        if keywords:
            doc.add_heading('Keywords', level=1)
            keywords_p = doc.add_paragraph()
            keywords_p.add_run(", ".join(keywords))
        
        # Footer
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run("Generated by DocuVision - Intelligent Document Processing System")
        footer_run.italic = True
        footer_run.font.size = Pt(9)
        
        doc.save(output_path)
        
        logger.info(f"Exported Word document to {output_path}")
        return output_path
    
    async def to_excel(self, result: Dict[str, Any], task_id: str) -> str:
        """
        Export tables to Excel file
        
        Args:
            result: Processing result dictionary
            task_id: Task identifier
        
        Returns:
            Path to exported XLSX file
        """
        try:
            import pandas as pd
        except ImportError:
            logger.error("pandas not installed, cannot export Excel")
            raise RuntimeError("pandas not installed")
        
        output_path = os.path.join(self.output_dir, f"{task_id}_tables.xlsx")
        
        tables = result.get('tables', [])
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Summary sheet
            summary_data = {
                'Property': ['File Name', 'Processed At', 'Total Pages', 'Total Tables'],
                'Value': [
                    result.get('document_info', {}).get('file_name', 'N/A'),
                    result.get('document_info', {}).get('processed_at', 'N/A'),
                    str(result.get('document_info', {}).get('pages', 'N/A')),
                    str(len(tables))
                ]
            }
            df_summary = pd.DataFrame(summary_data)
            df_summary.to_excel(writer, sheet_name='Summary', index=False)
            
            # Individual table sheets
            for idx, table in enumerate(tables):
                if 'data' in table and table['data']:
                    data = table['data']
                    
                    if len(data) > 1:
                        # First row as header
                        header = data[0]
                        num_cols = len(header)
                        
                        # 确保所有行都有相同的列数
                        normalized_data = []
                        for row in data[1:]:
                            # 如果行数少于列数，用空字符串填充
                            # 如果行数多于列数，截断
                            normalized_row = (row + [''] * num_cols)[:num_cols]
                            normalized_data.append(normalized_row)
                        
                        df = pd.DataFrame(normalized_data, columns=header)
                    else:
                        df = pd.DataFrame(data)
                    
                    sheet_name = f"Table_{idx + 1}"
                    df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=1)
                    worksheet = writer.sheets.get(sheet_name)
                    if worksheet is None:
                        worksheet = writer.book[sheet_name]
                    worksheet.cell(row=1, column=1, value=format_table_export_title(idx + 1, table))
        
        logger.info(f"Exported Excel to {output_path}")
        return output_path
    
    async def to_structured_json(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """
        Export to Azure Document Intelligence compatible JSON format
        
        Args:
            result: Processing result dictionary
        
        Returns:
            Structured JSON dictionary in Azure DI format
        """
        return {
            "apiVersion": "1.0.0",
            "modelId": "docuvision-general",
            "content": result.get('layout', {}).get('elements', []),
            "pages": [
                {
                    "pageNumber": page.get('page'),
                    "width": 0,
                    "height": 0,
                    "unit": "pixel",
                    "elements": [
                        e for e in result.get('layout', {}).get('elements', [])
                        if e.get('page') == page.get('page')
                    ]
                }
                for page in result.get('layout', {}).get('page_layouts', [])
            ],
            "tables": [
                {
                    "rowCount": t.get('rows', 0),
                    "columnCount": t.get('columns', 0),
                    "cells": self._convert_table_cells(t),
                    "boundingRegions": [{
                        "pageNumber": t.get('page', 1),
                        "polygon": []
                    }]
                }
                for t in result.get('tables', [])
            ],
            "keyValuePairs": result.get('extracted_fields', {}),
            "styles": [],
            "languages": ["en"],
            "contentFormat": "text"
        }
    
    def _convert_table_cells(self, table: Dict) -> List[Dict]:
        """Convert table data to cell format"""
        cells = []
        data = table.get('data', [])
        
        for row_idx, row in enumerate(data):
            for col_idx, cell_text in enumerate(row):
                cells.append({
                    "rowIndex": row_idx,
                    "columnIndex": col_idx,
                    "content": str(cell_text),
                    "kind": "columnHeader" if row_idx == 0 else "content"
                })
        
        return cells
